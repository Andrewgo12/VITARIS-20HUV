"""
Text Extraction Module for VITAL RED Gmail Integration
Hospital Universitaria ESE - Departamento de Innovación y Desarrollo
"""

import os
import io
import tempfile
from pathlib import Path
from typing import Optional, Dict, Any
import structlog

# Document processing imports
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract

# Optional imports for additional formats
try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

try:
    import cv2
    import numpy as np
    OPENCV_AVAILABLE = True
except ImportError:
    OPENCV_AVAILABLE = False

from config import FILE_CONFIG, PROCESSING_CONFIG

logger = structlog.get_logger(__name__)

class TextExtractor:
    """
    Comprehensive text extraction from various document formats
    """
    
    def __init__(self):
        self.logger = logger.bind(component="text_extractor")
        self.ocr_config = FILE_CONFIG["OCR_CONFIG"]
        
        # Configure Tesseract if available
        self._configure_tesseract()
    
    def _configure_tesseract(self):
        """Configure Tesseract OCR settings"""
        try:
            # Set Tesseract configuration
            self.tesseract_config = f'--oem {self.ocr_config["OEM"]} --psm {self.ocr_config["PSM"]}'
            
            # Test if Tesseract is available
            pytesseract.get_tesseract_version()
            self.tesseract_available = True
            self.logger.info("Tesseract OCR configured successfully")
            
        except Exception as e:
            self.tesseract_available = False
            self.logger.warning("Tesseract OCR not available", error=str(e))
    
    def extract_text(self, file_path: Path, mime_type: str) -> Optional[str]:
        """
        Main method to extract text from various file formats
        
        Args:
            file_path: Path to the file
            mime_type: MIME type of the file
            
        Returns:
            Extracted text or None if extraction failed
        """
        try:
            self.logger.info("Extracting text from file", 
                           file_path=str(file_path), mime_type=mime_type)
            
            if not file_path.exists():
                self.logger.error("File not found", file_path=str(file_path))
                return None
            
            # Route to appropriate extraction method based on MIME type
            if mime_type == "application/pdf":
                return self._extract_from_pdf(file_path)
            elif mime_type == "application/msword":
                return self._extract_from_doc(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx(file_path)
            elif mime_type.startswith("image/"):
                return self._extract_from_image(file_path)
            else:
                # Try generic text extraction
                return self._extract_generic(file_path)
                
        except Exception as e:
            self.logger.error("Text extraction failed", 
                            file_path=str(file_path), error=str(e))
            return None
    
    def _extract_from_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF files"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                self.logger.debug(f"PDF has {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                        else:
                            # If no text extracted, try OCR on the page
                            self.logger.debug(f"No text found on page {page_num + 1}, attempting OCR")
                            ocr_text = self._ocr_pdf_page(file_path, page_num)
                            if ocr_text:
                                text_content.append(ocr_text)
                                
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num + 1}", 
                                          error=str(e))
                        continue
            
            extracted_text = "\n\n".join(text_content)
            
            # If very little text was extracted, try OCR on the entire document
            if len(extracted_text.strip()) < 100:
                self.logger.info("Limited text extracted, attempting full OCR")
                ocr_text = self._ocr_entire_pdf(file_path)
                if ocr_text and len(ocr_text) > len(extracted_text):
                    extracted_text = ocr_text
            
            self.logger.info(f"Extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            self.logger.error("PDF text extraction failed", error=str(e))
            return None
    
    def _extract_from_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text_content)
            self.logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            self.logger.error("DOCX text extraction failed", error=str(e))
            return None
    
    def _extract_from_doc(self, file_path: Path) -> Optional[str]:
        """Extract text from DOC files using textract"""
        try:
            if not TEXTRACT_AVAILABLE:
                self.logger.warning("Textract not available for DOC extraction")
                return None
            
            text = textract.process(str(file_path)).decode('utf-8')
            self.logger.info(f"Extracted {len(text)} characters from DOC")
            return text
            
        except Exception as e:
            self.logger.error("DOC text extraction failed", error=str(e))
            return None
    
    def _extract_from_image(self, file_path: Path) -> Optional[str]:
        """Extract text from image files using OCR"""
        try:
            if not self.tesseract_available:
                self.logger.warning("Tesseract not available for image OCR")
                return None
            
            # Preprocess image for better OCR
            processed_image = self._preprocess_image(file_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(
                processed_image,
                lang=self.ocr_config["LANGUAGE"],
                config=self.tesseract_config
            )
            
            self.logger.info(f"Extracted {len(text)} characters from image via OCR")
            return text
            
        except Exception as e:
            self.logger.error("Image OCR failed", error=str(e))
            return None
    
    def _extract_generic(self, file_path: Path) -> Optional[str]:
        """Generic text extraction for unknown file types"""
        try:
            if TEXTRACT_AVAILABLE:
                text = textract.process(str(file_path)).decode('utf-8')
                self.logger.info(f"Extracted {len(text)} characters using textract")
                return text
            else:
                # Try reading as plain text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                    self.logger.info(f"Extracted {len(text)} characters as plain text")
                    return text
                    
        except Exception as e:
            self.logger.error("Generic text extraction failed", error=str(e))
            return None
    
    def _preprocess_image(self, file_path: Path) -> Image.Image:
        """Preprocess image for better OCR results"""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # If OpenCV is available, use advanced preprocessing
            if OPENCV_AVAILABLE:
                # Convert PIL to OpenCV format
                opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                
                # Convert to grayscale
                gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
                
                # Apply noise reduction
                denoised = cv2.medianBlur(gray, 3)
                
                # Apply adaptive thresholding
                thresh = cv2.adaptiveThreshold(
                    denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
                )
                
                # Convert back to PIL
                processed_image = Image.fromarray(thresh)
            else:
                # Basic preprocessing without OpenCV
                # Convert to grayscale
                processed_image = image.convert('L')
                
                # Enhance contrast
                from PIL import ImageEnhance
                enhancer = ImageEnhance.Contrast(processed_image)
                processed_image = enhancer.enhance(2.0)
            
            return processed_image
            
        except Exception as e:
            self.logger.warning("Image preprocessing failed, using original", error=str(e))
            return Image.open(file_path)
    
    def _ocr_pdf_page(self, file_path: Path, page_num: int) -> Optional[str]:
        """Perform OCR on a specific PDF page"""
        try:
            if not self.tesseract_available:
                return None
            
            # This would require converting PDF page to image first
            # For now, we'll return None and rely on the full PDF OCR
            return None
            
        except Exception as e:
            self.logger.error(f"OCR failed for PDF page {page_num}", error=str(e))
            return None
    
    def _ocr_entire_pdf(self, file_path: Path) -> Optional[str]:
        """Perform OCR on entire PDF by converting to images"""
        try:
            if not self.tesseract_available:
                return None
            
            # This would require pdf2image library
            # For now, we'll return None
            self.logger.info("Full PDF OCR not implemented yet")
            return None
            
        except Exception as e:
            self.logger.error("Full PDF OCR failed", error=str(e))
            return None
    
    def get_extraction_confidence(self, file_path: Path, mime_type: str) -> float:
        """
        Get confidence score for text extraction
        
        Returns:
            Confidence score between 0.0 and 1.0
        """
        try:
            if mime_type == "application/pdf":
                # For PDFs, confidence depends on whether text is embedded or OCR'd
                return 0.9  # High confidence for embedded text
            elif mime_type in ["application/msword", 
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                return 0.95  # Very high confidence for Word documents
            elif mime_type.startswith("image/"):
                # For images, confidence depends on OCR quality
                return 0.7  # Medium confidence for OCR
            else:
                return 0.5  # Low confidence for unknown formats
                
        except Exception:
            return 0.0
    
    def extract_metadata(self, file_path: Path, mime_type: str) -> Dict[str, Any]:
        """Extract metadata from files"""
        metadata = {
            'file_size': file_path.stat().st_size,
            'creation_time': file_path.stat().st_ctime,
            'modification_time': file_path.stat().st_mtime,
            'mime_type': mime_type
        }
        
        try:
            if mime_type == "application/pdf":
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        'page_count': len(pdf_reader.pages),
                        'pdf_info': pdf_reader.metadata if pdf_reader.metadata else {}
                    })
            
            elif mime_type.startswith("image/"):
                with Image.open(file_path) as img:
                    metadata.update({
                        'image_size': img.size,
                        'image_mode': img.mode,
                        'image_format': img.format
                    })
            
        except Exception as e:
            self.logger.warning("Failed to extract metadata", error=str(e))
        
        return metadata
    
    def is_supported_format(self, mime_type: str) -> bool:
        """Check if file format is supported for text extraction"""
        supported_formats = []
        for format_list in FILE_CONFIG["SUPPORTED_FORMATS"].values():
            supported_formats.extend(format_list)
        
        return mime_type in supported_formats
    
    def estimate_processing_time(self, file_path: Path, mime_type: str) -> float:
        """Estimate processing time in seconds"""
        try:
            file_size = file_path.stat().st_size
            
            # Base time estimates per MB
            time_per_mb = {
                "application/pdf": 2.0,  # 2 seconds per MB
                "application/msword": 1.0,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": 1.0,
                "image/": 5.0  # OCR is slower
            }
            
            # Get appropriate time estimate
            estimate = 1.0  # Default 1 second
            for format_prefix, time_est in time_per_mb.items():
                if mime_type.startswith(format_prefix):
                    estimate = time_est
                    break
            
            # Calculate based on file size
            size_mb = file_size / (1024 * 1024)
            estimated_time = size_mb * estimate
            
            # Minimum 1 second, maximum 300 seconds (5 minutes)
            return max(1.0, min(300.0, estimated_time))
            
        except Exception:
            return 30.0  # Default 30 seconds
